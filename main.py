from flask import Flask, request, render_template_string
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload

import os, io, time
import PyPDF2
import docx
import pandas as pd
import re


SCOPES = ['https://www.googleapis.com/auth/drive.file']
app = Flask(__name__)

def authenticate():
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    if not creds or not creds.valid:
        flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
        creds = flow.run_local_server(port=0)
        with open('token.json', 'w') as token:
            token.write(creds.to_json())
    return creds

def upload_file(file_path):
    creds = authenticate()
    service = build('drive', 'v3', credentials=creds)
    file_metadata = {'name': os.path.basename(file_path)}
    media = MediaFileUpload(file_path, resumable=True)
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    return file.get('id')

def download_file(file_id, destination_path):
    creds = authenticate()
    service = build('drive', 'v3', credentials=creds)
    request = service.files().get_media(fileId=file_id)
    with open(destination_path, 'wb') as f:
        downloader = MediaIoBaseDownload(f, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()

def extract_title_pdf(file_path):
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        if reader.metadata and '/Title' in reader.metadata:
            return reader.metadata['/Title']
        else:
            return reader.pages[0].extract_text().split('\n')[0]

def extract_title_word(file_path):
    doc = docx.Document(file_path)
    return doc.paragraphs[0].text if doc.paragraphs else "No Title"

def extract_text_pdf(file_path):
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        return "".join([page.extract_text() or "" for page in reader.pages])

def extract_text_word(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def search_and_highlight(text, keywords):
    matches = []
    for word in keywords:
        pattern = re.compile(re.escape(word.strip()), re.IGNORECASE)
        if re.search(pattern, text):
            matches.append(word.strip())
            text = pattern.sub(lambda m: f"<mark>{m.group(0)}</mark>", text)
    return matches, text

def classify(text):
    text = text.lower()
    scientific_terms = ['science', 'research', 'experiment', 'data', 'cloud', 'algorithm', 'analysis', 'ai', 'machine learning']
    religious_terms = ['religion', 'quran', 'islam', 'hadith', 'prayer']
    literary_terms = ['poetry', 'novel', 'story', 'literature', 'poem']
    if any(word in text for word in religious_terms):
        return "ديني"
    elif any(word in text for word in scientific_terms):
        return "علمي"
    elif any(word in text for word in literary_terms):
        return "أدبي"
    else:
        return "عام"

 
html_form = """
<!doctype html>
<title>تحليل المستندات السحابي</title>
<link rel="stylesheet" href="{{ url_for('static', filename='style.css') }}">

<h2>📤 رفع وتحليل مستند (PDF أو Word)</h2>
<form method=post enctype=multipart/form-data>
  <input type=file name=file required><br><br>
  كلمات مفتاحية  :<br>
  <input type=text name=keywords><br><br>
  <input type=submit value="رفع وتحليل">
</form>
<br>
<a href="/history">📁 عرض كل الملفات التي تم تحليلها</a><br>
<a href="/sorted">🗂️ عرض الملفات مرتبة حسب العنوان</a>
<a href="/re-search">🔁 إعادة البحث في ملف سابق</a><br><br>

{% if result %}
  <h3>🔍 النتائج:</h3>
  <pre>{{ result }}</pre>
  {% if highlighted %}
  <h4>📄 مقتطف من النص بعد التمييز:</h4>
  <div style="background:#f4f4f4;padding:10px;border:1px solid #ccc;max-height:300px;overflow:auto;">
    {{ highlighted|safe }}
  </div>
  {% endif %}
{% endif %}
"""

@app.route("/", methods=['GET', 'POST'])
def index():
    result = ""
    highlighted = ""
    if request.method == 'POST':
        start_time = time.time()
        file = request.files['file']
        keywords = request.form['keywords'].split(',') if request.form['keywords'] else []
        filename = file.filename.lower()
        os.makedirs("temp", exist_ok=True)
        filepath = os.path.join("temp", filename)
        file.save(filepath)

        file_id = upload_file(filepath)
        download_path = os.path.join("temp", "downloaded_" + filename)
        download_file(file_id, download_path)

        if filename.endswith(".pdf"):
            text = extract_text_pdf(download_path)
            title = extract_title_pdf(download_path)
        elif filename.endswith(".docx"):
            text = extract_text_word(download_path)
            title = extract_title_word(download_path)
        else:
            return render_template_string(html_form, result="❌ صيغة غير مدعومة.", highlighted=None)

        matches, highlighted_text = search_and_highlight(text, keywords)
        category = classify(text)
        size_kb = os.path.getsize(download_path) / 1024
        elapsed_time = time.time() - start_time

        result = f"""📄 العنوان: {title}
📚 التصنيف: {category}
🔍 الكلمات المطابقة: {matches}
📦 الحجم: {size_kb:.2f} KB
⏱️ زمن المعالجة: {elapsed_time:.2f} ثانية
🆔 File ID: {file_id}"""

        data = {
            "File Name": filename,
            "Title": title,
            "Classification": category,
            "Found Keywords": ", ".join(matches),
            "Size (KB)": size_kb,
            "Processing Time (s)": elapsed_time,
            "Drive File ID": file_id
        }

        if os.path.exists("results.xlsx"):
            df_existing = pd.read_excel("results.xlsx")
            df_new = pd.DataFrame([data])
            df_combined = pd.concat([df_existing, df_new], ignore_index=True)
            df_combined.to_excel("results.xlsx", index=False)
        else:
            pd.DataFrame([data]).to_excel("results.xlsx", index=False)

        return render_template_string(html_form, result=result, highlighted=highlighted_text)

    return render_template_string(html_form, result=None, highlighted=None)

@app.route("/history")
def history():
    if not os.path.exists("results.xlsx"):
        return "<h3> لا توجد نتائج محفوظة بعد.</h3>"
    df = pd.read_excel("results.xlsx")
    return f"""<h2>📁 الملفات التي تم تحليلها</h2>
    <a href="/">⬅️ الرجوع</a><br><br>{df.to_html(index=False)}"""

@app.route("/sorted")
def sorted_docs():
    if not os.path.exists("results.xlsx"):
        return "<h3> لا توجد نتائج محفوظة بعد.</h3>"
    df = pd.read_excel("results.xlsx")
    df = df.sort_values(by="Title")
    return f"""<h2>🗂️ الملفات مرتبة حسب العنوان</h2>
    <a href="/">⬅️ الرجوع</a><br><br>{df.to_html(index=False)}"""

@app.route("/re-search", methods=["GET", "POST"])
def re_search():
    result = ""
    highlighted = ""
    if not os.path.exists("results.xlsx"):
        return "<h3> لا توجد ملفات محفوظة.</h3>"

    df = pd.read_excel("results.xlsx")

    if request.method == "POST":
        file_id = request.form["file_id"]
        keywords = request.form["keywords"].split(',')

        selected_row = df[df["Drive File ID"] == file_id]
        if not selected_row.empty:
            original_filename = selected_row["File Name"].values[0]
            ext = os.path.splitext(original_filename)[1]
            filename = "temp/reanalyzed_file" + ext
        else:
            return "❌ ملف غير موجود."

        download_file(file_id, filename)

        if filename.endswith(".pdf"):
            text = extract_text_pdf(filename)
        elif filename.endswith(".docx"):
            text = extract_text_word(filename)
        else:
            return "❌ صيغة غير مدعومة"

        matches, highlighted_text = search_and_highlight(text, keywords)
        result = f"🔍 الكلمات المطابقة: {matches}"

        return render_template_string(html_form, result=result, highlighted=highlighted_text)

    options = "".join(
        [f"<option value='{row['Drive File ID']}'>{row['Title']} ({row['File Name']})</option>" for _, row in df.iterrows()]
    )
    form_html = f"""
    <h2>🔁 إعادة البحث</h2>
    <form method='post'>
    اختر الملف:
    <select name='file_id'>{options}</select><br><br>
    كلمات البحث: <input name='keywords'><br><br>
    <input type='submit' value='بحث'>
    </form><br><a href='/'>⬅️ الرجوع</a>"""

    return form_html

if __name__ == "__main__":
    app.run(debug=True)
